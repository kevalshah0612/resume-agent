"""
resume.py — Simple Resume Builder for Keval Shah
================================================

Only 3 commands:

1. JSON → DOCX
   python resume.py temp.json
   python resume.py temp.json "Palo Alto Networks"
   python resume.py temp.json 6           # section gap = 6pts
   python resume.py temp.json 6 4         # section gap = 6pts, sub-section gap = 4pts
   python resume.py temp.json "Palo Alto Networks" 6
   python resume.py temp.json "Palo Alto Networks" 6 4

2. DOCX → PDF + archive + delete DOCX
   python resume.py resume.docx
   python resume.py resume.docx "Palo Alto Networks"

   This keeps DOCX files in Resume-word, keeps final PDFs in Resume-pdf,
   copies one PDF into Resume-pdf/archives/YYYY-MM-DD/, and deletes the DOCX.

3. Clear generated DOCX/PDF files
   python resume.py clear

Gap args (only for JSON → DOCX):
  - One numeric arg  → section gap in pts   (default: 4)
  - Two numeric args → section gap, then sub-section gap in pts   (default: 4, 4)
  Section gap    = space added after each rendered section block.
  Sub-section gap = space added between entries inside a section
                    (e.g. between two jobs, or between two projects).

Rules:
  - config.type controls resume type: backend, fullstack, aiml, aitool
  - config.level controls base level: 2=Entry/SWE I, 3=Mid, 4=Intern
  - config.layout_profile controls section order and experience order when present
  - Skills render near the top for cold-apply layouts
  - If layout_profile is missing, level 3 renders TCS first and level 2/4 renders internship first
  - Python never changes role titles. It renders titles exactly from JSON.
  - Company passed in command overrides company/output filename only.
  - PDF conversion tries Microsoft Word first, then LibreOffice fallback.
  - Adds space after each rendered section, but never after the final rendered section.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app_properties import ARCHIVES_DIR_NAME, RESUME_STEM, WORD_DIR_NAME, PDF_DIR_NAME
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# ── Formatting constants ─────────────────────────────────────────────────────
DEFAULT_FONT = "Arial"
BODY_PT = Pt(10.5)
SUB_PT = Pt(10.5)
SEC_PT = Pt(10.5)
NAME_PT = Pt(14)
CONTACT_PT = Pt(12)
ZERO = Pt(0)

PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
M_TOP = Inches(0.50)
M_BOT = Inches(0.94)
M_LEFT = Inches(1.00)
M_RIGHT = Inches(1.00)
TAB_PAD = Inches(0.05)
LINE_SPACING = 360          # 360 = 1.5 line spacing in Word XML
BULLET_LEFT_TWIPS = 720     # 0.5 inch
BULLET_HANGING_TWIPS = 360  # 0.25 inch
INDENT_LEFT = Inches(0.5)

DEFAULT_SECTION_GAP = 4   # pts between sections
DEFAULT_SUB_GAP = 4       # pts between entries within a section

ARCHIVES_DIR = ARCHIVES_DIR_NAME
WORD_DIR = WORD_DIR_NAME
PDF_DIR = PDF_DIR_NAME
TEMPLATE_DOCX = "Headless+Resume+Template.docx"


# ── Configs ──────────────────────────────────────────────────────────────────
CONFIGS = {
    ("backend", 2): {"label": "SWE Backend Entry", "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("backend", 3): {"label": "SWE Backend Mid",   "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("backend", 4): {"label": "SWE Backend Intern", "grad": "Expected Aug 2026", "order": ["education", "skills", "experience", "projects"]},
    ("fullstack", 2): {"label": "Fullstack Entry",  "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("fullstack", 3): {"label": "Fullstack Mid",    "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("fullstack", 4): {"label": "Fullstack Intern", "grad": "Expected Aug 2026", "order": ["education", "skills", "experience", "projects"]},
    ("aiml", 2): {"label": "AI/ML Entry",           "grad": "Expected Aug 2026", "order": ["education", "skills", "projects", "experience"]},
    ("aiml", 3): {"label": "AI/ML Mid",             "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("aiml", 4): {"label": "AI/ML Intern",          "grad": "Expected Aug 2026", "order": ["education", "skills", "experience", "projects"]},
    ("aitool", 2): {"label": "AI Tooling Entry",    "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("aitool", 3): {"label": "AI Tooling Mid",      "grad": "Expected Aug 2026", "order": ["summary", "skills", "experience", "projects", "education"]},
    ("aitool", 4): {"label": "AI Tooling Intern",   "grad": "Expected Aug 2026", "order": ["education", "skills", "experience", "projects"]},
}


LAYOUT_ORDERS = {
    "student_entry":       ["education", "skills", "experience", "projects"],
    "professional_entry":  ["summary", "skills", "experience", "projects", "education"],
    "mid":                 ["summary", "skills", "experience", "projects", "education"],
    "aiml_entry":          ["education", "skills", "projects", "experience"],
    "aitool_mid":          ["summary", "skills", "experience", "projects", "education"],
    "internship":          ["education", "skills", "experience", "projects"],
}

GHI_FIRST_LAYOUTS = {"student_entry", "aiml_entry", "internship"}
TCS_FIRST_LAYOUTS = {"professional_entry", "mid", "aitool_mid"}

ATS_KEYWORDS = {
    "backend": "Software Engineer, Backend Engineer, Java, Python, Spring Boot, AWS, Microservices, Docker, REST APIs, Distributed Systems, PostgreSQL, Redis, Kubernetes",
    "fullstack": "Software Engineer, Full Stack, Java, Python, React, TypeScript, JavaScript, Spring Boot, AWS, Microservices, Docker, REST APIs, PostgreSQL",
    "aiml": "AI Engineer, Machine Learning, LLM, RAG, LangChain, Python, FastAPI, MLOps, Docker, AWS, pgvector, HuggingFace, model evaluation",
    "aitool": "Software Engineer, AI Tooling, Claude Code, AI agents, Python, Bash, Linux, automation, scripting, developer workflows, Docker, AWS, GitHub Actions",
}


# ── Basic helpers ────────────────────────────────────────────────────────────
def clean(value: Any) -> str:
    text = "" if value is None else str(value)
    for ch in ["\u200B", "\u200C", "\u200D", "\u2060", "\uFEFF", "\u180E"]:
        text = text.replace(ch, "")
    return text.strip()


def safe_name(value: str) -> str:
    return re.sub(r"[^\w]+", "_", clean(value)).strip("_") or "Company"


def fail(message: str) -> None:
    print(f"Error: {message}")
    sys.exit(1)


def load_json(path: str) -> Dict[str, Any]:
    if not os.path.exists(path):
        fail(f"JSON file not found: {path}")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except json.JSONDecodeError as e:
        fail(f"Invalid JSON: {e}")


def parse_extra_args(args: list[str]) -> tuple[str, int, int]:
    """
    Parse everything after the first positional arg (json/docx path).

    Returns (company, section_gap, sub_gap).

    Accepted patterns (args = sys.argv[2:]):
      []                        -> ("", 4, 4)
      ["6"]                     -> ("", 6, 4)
      ["6", "4"]                -> ("", 6, 4)
      ["Palo Alto Networks"]    -> ("Palo Alto Networks", 4, 4)
      ["Palo Alto Networks", "6"]      -> ("Palo Alto Networks", 6, 4)
      ["Palo Alto Networks", "6", "4"] -> ("Palo Alto Networks", 6, 4)

    Numeric-only tokens at the END of the arg list are treated as gap values.
    Everything before those is joined as the company name.
    """
    section_gap = DEFAULT_SECTION_GAP
    sub_gap = DEFAULT_SUB_GAP
    company = ""

    extras = list(args)

    # Pull up to 2 trailing numeric tokens as gap values.
    gap_vals: list[int] = []
    while extras and extras[-1].isdigit() and len(gap_vals) < 2:
        gap_vals.insert(0, int(extras.pop()))

    if len(gap_vals) >= 1:
        section_gap = gap_vals[0]
    if len(gap_vals) >= 2:
        sub_gap = gap_vals[1]

    company = clean(" ".join(extras))
    return company, section_gap, sub_gap


def normalize_type(value: Any) -> str:
    raw = clean(value).lower().replace("-", "").replace("_", "").replace(" ", "")
    aliases = {
        "backend": "backend",
        "swebackend": "backend",
        "fullstack": "fullstack",
        "full": "fullstack",
        "aiml": "aiml",
        "ai": "aiml",
        "ml": "aiml",
        "aitool": "aitool",
        "aitooling": "aitool",
        "devproductivity": "aitool",
    }
    return aliases.get(raw, "fullstack")


def normalize_level(value: Any) -> int:
    raw = clean(value).lower().replace("-", "").replace("_", "").replace(" ", "")
    aliases = {
        "1": 4,
        "2": 2,
        "3": 3,
        "4": 4,
        "entry": 2,
        "newgrad": 2,
        "junior": 2,
        "swe1": 2,
        "mid": 3,
        "midlevel": 3,
        "swe2": 3,
        "intern": 4,
        "internship": 4,
        "coop": 4,
    }
    return aliases.get(raw, 3)


def clean_bullet_text(text: Any) -> str:
    text = clean(text)
    while text.endswith((".", ";", ":")):
        text = text[:-1].rstrip()
    return text


def normalize_layout_profile(value: Any, rtype: str, level: int) -> str:
    raw = clean(value).lower().replace("-", "_").replace(" ", "_")

    aliases = {
        "student": "student_entry",
        "student_entry": "student_entry",
        "entry_student": "student_entry",
        "professional": "professional_entry",
        "professional_entry": "professional_entry",
        "entry_professional": "professional_entry",
        "swe1_professional": "professional_entry",
        "mid": "mid",
        "mid_level": "mid",
        "aiml_entry": "aiml_entry",
        "ai_entry": "aiml_entry",
        "ml_entry": "aiml_entry",
        "aiml_mid_product": "mid",
        "ai_mid_product": "mid",
        "llm_product": "mid",
        "rag_product": "mid",
        "aiml_mid_platform": "mid",
        "ai_mid_platform": "mid",
        "ml_platform": "mid",
        "aitool_mid": "aitool_mid",
        "ai_tooling_mid": "aitool_mid",
        "developer_productivity": "aitool_mid",
        "intern": "internship",
        "internship": "internship",
        "coop": "internship",
    }

    if raw in aliases:
        return aliases[raw]

    if level == 4:
        return "internship"
    if level == 3:
        if rtype == "aiml":
            return "mid"
        if rtype == "aitool":
            return "aitool_mid"
        return "mid"
    return "student_entry"


def config(data: dict) -> dict:
    return data.get("config") or {}


def get_jobs(data: dict) -> list:
    return data.get("experience") or data.get("professional_experience") or []


def get_projects(data: dict) -> list:
    return data.get("projects") or []


def get_skills_rows(data: dict) -> list[tuple[str, str]]:
    if isinstance(data.get("skills"), dict):
        rows = []
        s = data["skills"]
        for i in range(1, 8):
            label = clean(s.get(f"row{i}_label", ""))
            terms = s.get(f"row{i}_terms", [])
            if isinstance(terms, list):
                terms = ", ".join(clean(t) for t in terms if clean(t))
            terms = clean(terms)
            if label and terms:
                rows.append((label, terms))
        return rows

    skills = data.get("technical_skills") or {}
    return [(clean(k), clean(v)) for k, v in skills.items() if clean(k) and clean(v)]


def get_project_tech(project: dict) -> str:
    tech = project.get("tech_label") or project.get("tech", "")
    if isinstance(tech, list):
        return ", ".join(clean(t) for t in tech if clean(t))
    return clean(tech)


def ordered_experience(data: dict, level: int, layout_profile: str = "") -> list:
    jobs = get_jobs(data)
    ghi = [j for j in jobs if "global health impact" in clean(j.get("company", "")).lower()]
    others = [j for j in jobs if "global health impact" not in clean(j.get("company", "")).lower()]
    order = clean(config(data).get("experience_order", "")).lower().replace("-", "_").replace(" ", "_")

    if order in {"json", "json_order"}:
        return jobs
    if order in {"ghi", "ghi_first", "internship_first"}:
        return ghi + others
    if order in {"tcs", "tcs_first", "professional_first"}:
        return others + ghi

    if layout_profile in TCS_FIRST_LAYOUTS:
        return others + ghi
    if layout_profile in GHI_FIRST_LAYOUTS:
        return ghi + others

    return others + ghi if level == 3 else ghi + others


def experience_order_label(data: dict, level: int, layout_profile: str = "") -> str:
    order = clean(config(data).get("experience_order", "")).lower().replace("-", "_").replace(" ", "_")

    if order in {"json", "json_order"}:
        return "JSON order"
    if order in {"ghi", "ghi_first", "internship_first"}:
        return "GHI first"
    if order in {"tcs", "tcs_first", "professional_first"}:
        return "TCS first"
    if layout_profile in TCS_FIRST_LAYOUTS:
        return "TCS first"
    if layout_profile in GHI_FIRST_LAYOUTS:
        return "Internship/GHI first"
    return "TCS first" if level == 3 else "Internship/GHI first"


def bool_value(value: Any, default: bool = False) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return default
    raw = clean(value).lower()
    if raw in {"true", "yes", "1", "y"}:
        return True
    if raw in {"false", "no", "0", "n"}:
        return False
    return default


def build_render_profile(data: dict) -> dict[str, Any]:
    """Describe the exact content order the DOCX renderer will use."""
    cfg = config(data)
    rtype = normalize_type(cfg.get("type", "fullstack"))
    level = normalize_level(cfg.get("level", 3))
    layout_profile = normalize_layout_profile(cfg.get("layout_profile", ""), rtype, level)
    conf = CONFIGS[(rtype, level)]
    configured_sections = LAYOUT_ORDERS.get(layout_profile, conf["order"])

    present = {
        "summary": bool(clean(data.get("summary", ""))),
        "education": bool(data.get("education")),
        "skills": bool(get_skills_rows(data)),
        "experience": bool(get_jobs(data)),
        "projects": bool(get_projects(data)),
    }
    rendered_sections = [section for section in configured_sections if present.get(section, False)]
    jobs = ordered_experience(data, level, layout_profile)
    projects = get_projects(data)
    education = data.get("education") or []
    render_ta = level in {2, 4} and bool_value(cfg.get("ta_active"), default=False)

    return {
        "resume_type": rtype,
        "level": level,
        "level_label": "Entry/SWE I" if level == 2 else "Mid" if level == 3 else "Intern",
        "layout_profile": layout_profile,
        "configured_section_order": list(configured_sections),
        "rendered_section_order": rendered_sections,
        "experience_order": [
            {
                "position": index,
                "company": clean(job.get("company", "")),
                "title": clean(job.get("title", "")),
                "location": clean(job.get("location", "")),
                "dates": clean(job.get("dates", "")),
                "employment_note": clean(job.get("employment_note", "")),
                "bullet_count": len(job.get("bullets") or []),
            }
            for index, job in enumerate(jobs, start=1)
        ],
        "education_order": [
            {
                "position": index,
                "university": clean(item.get("university", "")),
                "degree": clean(item.get("degree", "")),
                "graduation": clean(item.get("graduation", "")) or (conf["grad"] if index == 1 else ""),
            }
            for index, item in enumerate(education, start=1)
        ],
        "project_order": [
            {
                "position": index,
                "name": clean(project.get("name", "")),
                "technology": get_project_tech(project),
                "bullet_count": len(project.get("bullets") or []),
            }
            for index, project in enumerate(projects, start=1)
        ],
        "skill_rows": [
            {"position": index, "label": label, "terms": terms}
            for index, (label, terms) in enumerate(get_skills_rows(data), start=1)
        ],
        "ta_bullet_rendered_under_education": bool(
            render_ta
            and education
            and clean(education[0].get("ta_bullet", ""))
        ),
        "renderer_constraints": {
            "experience_order_is_calculated_from_layout": not bool(clean(cfg.get("experience_order", ""))),
            "empty_sections_are_skipped": True,
            "titles_render_exactly_from_json": True,
            "bullet_counts_must_remain_unchanged_during_final_review": True,
        },
    }


# ── PDF helpers ──────────────────────────────────────────────────────────────
def pdf_with_word(docx_path: str) -> str:
    try:
        import win32com.client
        import pywintypes
    except ImportError as e:
        raise RuntimeError("pywin32 not installed") from e

    abs_docx = os.path.abspath(docx_path)
    pdf_path = os.path.splitext(abs_docx)[0] + ".pdf"
    word = None
    doc = None

    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(abs_docx, False, True)
        doc.ExportAsFixedFormat(
            OutputFileName=pdf_path,
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            Range=0,
            Item=0,
            IncludeDocProps=True,
            KeepIRM=True,
            CreateBookmarks=0,
            DocStructureTags=True,
            BitmapMissingFonts=True,
            UseISO19005_1=False,
        )
    except pywintypes.com_error as e:
        raise RuntimeError(f"Word PDF export failed: {e}") from e
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass

    if not os.path.exists(pdf_path):
        raise RuntimeError("Word ran but PDF was not created")
    return pdf_path


def pdf_with_libreoffice(docx_path: str) -> str:
    abs_docx = os.path.abspath(docx_path)
    out_dir = os.path.dirname(abs_docx)
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, abs_docx],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        raise RuntimeError(result.stderr.strip() or result.stdout.strip() or "LibreOffice failed")

    pdf_path = os.path.splitext(abs_docx)[0] + ".pdf"
    if not os.path.exists(pdf_path):
        raise RuntimeError("LibreOffice ran but PDF was not created")
    return pdf_path


def infer_company_from_docx(docx_path: str) -> str:
    stem = os.path.splitext(os.path.basename(docx_path))[0]
    match = re.match(r"Keval_Shah_(.+?)_Resume$", stem, re.IGNORECASE)
    return match.group(1) if match else "Company"


def docx_to_pdf(docx_path: str, company_override: str = "") -> None:
    if not os.path.exists(docx_path):
        fail(f"DOCX not found: {docx_path}")

    company = company_override or infer_company_from_docx(docx_path)
    company_safe = safe_name(company)
    source_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
    pdf_dir = os.path.abspath(PDF_DIR)
    os.makedirs(pdf_dir, exist_ok=True)
    now = datetime.now()

    try:
        pdf_path = pdf_with_word(docx_path)
        converter = "Microsoft Word"
    except Exception as word_error:
        try:
            pdf_path = pdf_with_libreoffice(docx_path)
            converter = "LibreOffice"
        except Exception as libre_error:
            fail(f"Could not convert DOCX to PDF. Word error: {word_error}. LibreOffice error: {libre_error}")

    final_pdf = os.path.join(pdf_dir, f"{RESUME_STEM}_{company_safe}_Resume.pdf")

    if os.path.abspath(pdf_path) != os.path.abspath(final_pdf):
        if os.path.exists(final_pdf):
            os.remove(final_pdf)
        shutil.move(pdf_path, final_pdf)

    archive_dir = os.path.join(pdf_dir, ARCHIVES_DIR, now.strftime("%Y-%m-%d"))
    os.makedirs(archive_dir, exist_ok=True)

    archive_pdf = os.path.join(archive_dir, f"{RESUME_STEM}_{company_safe}_{now.strftime('%H-%M-%S')}.pdf")
    shutil.copy2(final_pdf, archive_pdf)

    try:
        os.remove(docx_path)
    except Exception as e:
        print(f"Warning: could not delete DOCX: {e}")

    lock_name = "~$" + os.path.basename(docx_path)[2:]
    lock_path = os.path.join(source_dir, lock_name)
    if os.path.exists(lock_path):
        try:
            os.remove(lock_path)
        except Exception:
            pass

    print(f"PDF saved: {final_pdf}")
    print(f"Converter: {converter}")
    print(f"Archived : {archive_pdf}")
    print(f"Deleted  : {docx_path}")


def clear_folder() -> None:
    deleted = 0
    for folder in [".", WORD_DIR, PDF_DIR]:
        if not os.path.isdir(folder):
            continue
        for name in os.listdir(folder):
            path = os.path.join(folder, name)
            if os.path.isdir(path):
                continue
            if name.lower().endswith((".docx", ".pdf")) or name.startswith("~$"):
                try:
                    os.remove(path)
                    deleted += 1
                    print(f"Deleted: {path}")
                except Exception as e:
                    print(f"Could not delete {path}: {e}")
    print(f"Done. Removed {deleted} file(s).")


# ── DOCX low level helpers ───────────────────────────────────────────────────
def do_layout(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width = PAGE_W
    s.page_height = PAGE_H
    s.top_margin = M_TOP
    s.bottom_margin = M_BOT
    s.left_margin = M_LEFT
    s.right_margin = M_RIGHT


def document_from_template() -> Document:
    """Start from the Headless template when present so Word styles/numbering match."""
    doc = Document(TEMPLATE_DOCX) if os.path.exists(TEMPLATE_DOCX) else Document()
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    do_layout(doc)
    return doc


def rtab(doc: Document) -> int:
    s = doc.sections[0]
    return int(s.page_width - s.left_margin - s.right_margin - TAB_PAD)


def rf(run, *, sz=BODY_PT, bold=False, italic=False) -> None:
    run.font.name = DEFAULT_FONT
    run.font.size = sz
    run.bold = bold
    run.italic = italic


def sp(paragraph, *, before=0, after=0, line=LINE_SPACING) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    for child in list(ppr):
        if child.tag == qn("w:spacing"):
            ppr.remove(child)
    el = OxmlElement("w:spacing")
    el.set(qn("w:before"), str(before))
    el.set(qn("w:after"), str(after))
    el.set(qn("w:line"), str(line))
    el.set(qn("w:lineRule"), "auto")
    ppr.append(el)


def headless_sp(paragraph, *, before=0, after=0, line=LINE_SPACING) -> None:
    sp(paragraph, before=before, after=after, line=line)


def gap(doc: Document, pts: int = 4) -> None:
    p = doc.add_paragraph()
    ppr = p._element.get_or_add_pPr()
    el = OxmlElement("w:spacing")
    el.set(qn("w:before"), "0")
    el.set(qn("w:after"), "0")
    el.set(qn("w:line"), str(pts * 20))
    el.set(qn("w:lineRule"), "exact")
    ppr.append(el)


def remove_final_empty_gap(doc: Document) -> None:
    """Remove trailing empty paragraphs at the end of the document."""
    while doc.paragraphs:
        last = doc.paragraphs[-1]
        if clean(last.text):
            break
        element = last._element
        element.getparent().remove(element)


def force_bold(run) -> None:
    rpr = run._r.get_or_add_rPr()
    b = OxmlElement("w:b")
    b.set(qn("w:val"), "1")
    rpr.append(b)
    bcs = OxmlElement("w:bCs")
    bcs.set(qn("w:val"), "1")
    rpr.append(bcs)


def hyperlink(paragraph, text: str, url: str, size=CONTACT_PT) -> None:
    text = clean(text)
    url = clean(url)
    if not url:
        r = paragraph.add_run(text)
        rf(r, sz=size)
        return

    rid = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), DEFAULT_FONT)
    fonts.set(qn("w:hAnsi"), DEFAULT_FONT)
    rpr.append(fonts)

    size_el = OxmlElement("w:sz")
    size_el.set(qn("w:val"), str(int(size.pt * 2)))
    rpr.append(size_el)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    link.append(run)
    paragraph._p.append(link)


def star_bold(paragraph, text: str, bold_markers: bool = True) -> None:
    for part in re.split(r"(\*\*.*?\*\*)", clean(text)):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            rf(r)
            if bold_markers:
                force_bold(r)
        else:
            r = paragraph.add_run(part)
            rf(r)


# ── DOCX render blocks ───────────────────────────────────────────────────────
def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    headless_sp(p)
    r = p.add_run(clean(text))
    rf(r, sz=SEC_PT, bold=True)


def company_header(doc: Document, company: str, title: str, location: str, dates: str) -> None:
    p = doc.add_paragraph()
    headless_sp(p)
    tabs = p.paragraph_format.tab_stops
    tabs.clear_all()
    tabs.add_tab_stop(rtab(doc), WD_TAB_ALIGNMENT.RIGHT)

    label_parts = []
    if clean(title) and clean(company):
        label_parts.append(f"{clean(title)} at {clean(company)}")
    elif clean(title):
        label_parts.append(clean(title))
    elif clean(company):
        label_parts.append(clean(company))
    if clean(location):
        label_parts.append(clean(location))

    r = p.add_run(", ".join(label_parts))
    rf(r, sz=SUB_PT, italic=True)

    if clean(dates):
        tab = p.add_run("\t")
        rf(tab, sz=SUB_PT)
        r3 = p.add_run(clean(dates))
        rf(r3, sz=SUB_PT, italic=True)


def project_header(doc: Document, name: str, tech: str, url: str) -> None:
    p = doc.add_paragraph()
    headless_sp(p)
    tabs = p.paragraph_format.tab_stops
    tabs.clear_all()
    tabs.add_tab_stop(rtab(doc), WD_TAB_ALIGNMENT.RIGHT)

    r = p.add_run(clean(name))
    rf(r, sz=SUB_PT, italic=True)

    if clean(url):
        tab = p.add_run("\t")
        rf(tab, sz=SUB_PT)
        hyperlink(p, "Project Link", clean(url), size=SUB_PT)


def set_template_bullet(
    paragraph,
    *,
    num_id: int = 2,
    left_twips: int = BULLET_LEFT_TWIPS,
    hanging_twips: int = BULLET_HANGING_TWIPS,
) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    for child in list(ppr):
        if child.tag in {qn("w:numPr"), qn("w:ind")}:
            ppr.remove(child)

    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_el)
    ppr.append(num_pr)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    ppr.append(ind)


def bullet_paragraph(doc: Document, *, num_id: int = 2, line: int = LINE_SPACING):
    p = doc.add_paragraph()
    headless_sp(p, line=line)
    set_template_bullet(p, num_id=num_id)
    return p


def bullet(doc: Document, text: str, bold_markers: bool = True) -> None:
    text = clean_bullet_text(text)
    if not text:
        return
    p = bullet_paragraph(doc)
    star_bold(p, text, bold_markers)


# ── Section renderers ────────────────────────────────────────────────────────
def render_summary(doc: Document, data: dict, bold_markers: bool, sub_gap: int) -> bool:
    summary = clean(data.get("summary", ""))
    if not summary:
        return False
    section_heading(doc, "Summary")
    p = doc.add_paragraph()
    headless_sp(p)
    star_bold(p, summary, bold_markers)
    return True


def render_education(doc: Document, data: dict, grad: str, level: int, bold_markers: bool, sub_gap: int) -> bool:
    education = data.get("education") or []
    if not education:
        return False

    cfg = config(data)
    render_ta = level in {2, 4} and bool_value(cfg.get("ta_active"), default=False)

    section_heading(doc, "Education")

    for i, edu in enumerate(education):
        p = bullet_paragraph(doc, num_id=1)
        headless_sp(p)
        tabs = p.paragraph_format.tab_stops
        tabs.clear_all()
        tabs.add_tab_stop(rtab(doc), WD_TAB_ALIGNMENT.RIGHT)

        r = p.add_run(clean(edu.get("degree", "")))
        rf(r, sz=SUB_PT)

        # Prefer the graduation value from JSON so Prompt/Story remain source of truth.
        # Fall back to config default only when JSON omits the field.
        display_grad = clean(edu.get("graduation", "")) or (grad if i == 0 else "")
        if display_grad:
            if not display_grad.lower().startswith("status"):
                display_grad = f"{display_grad}"
            tab = p.add_run("\t")
            rf(tab, sz=SUB_PT)
            r3 = p.add_run(display_grad)
            rf(r3, sz=SUB_PT)

        university = clean(edu.get("university", ""))
        location = clean(edu.get("location", ""))
        if university or location:
            p2 = doc.add_paragraph()
            headless_sp(p2)
            p2.paragraph_format.left_indent = INDENT_LEFT
            tabs2 = p2.paragraph_format.tab_stops
            tabs2.clear_all()
            tabs2.add_tab_stop(rtab(doc), WD_TAB_ALIGNMENT.RIGHT)

            r4 = p2.add_run(university)
            rf(r4, sz=SUB_PT)
            if location:
                tab2 = p2.add_run("\t")
                rf(tab2, sz=SUB_PT)
                r5 = p2.add_run(location)
                rf(r5, sz=SUB_PT)

        if i == 0 and render_ta and clean(edu.get("ta_bullet", "")):
            bullet(doc, edu.get("ta_bullet", ""), bold_markers)

        if i < len(education) - 1:
            gap(doc, sub_gap)

    return True


def render_skills(doc: Document, data: dict, bold_markers: bool, sub_gap: int) -> bool:
    rows = get_skills_rows(data)
    if not rows:
        return False

    section_heading(doc, "Skills")
    for label, terms in rows:
        p = doc.add_paragraph()
        headless_sp(p)
        r = p.add_run(f"{label}: ")
        rf(r)
        star_bold(p, terms, bold_markers)
    return True


def render_experience(doc: Document, data: dict, level: int, layout_profile: str, bold_markers: bool, sub_gap: int) -> bool:
    jobs = ordered_experience(data, level, layout_profile)
    if not jobs:
        return False

    section_heading(doc, "Work History")
    for i, job in enumerate(jobs):
        company_header(doc, job.get("company", ""), job.get("title", ""), job.get("location", ""), job.get("dates", ""))
        employment_note = str(job.get("employment_note", "") or "").strip()

        if employment_note:
            note_p = doc.add_paragraph()
            headless_sp(note_p)

            note_run = note_p.add_run(employment_note)
            note_run.italic = True
            note_run.font.size = BODY_PT

        for b in job.get("bullets") or []:
            bullet(doc, b, bold_markers)
        if i < len(jobs) - 1:
            gap(doc, sub_gap)

    return True


def render_projects(doc: Document, data: dict, bold_markers: bool, sub_gap: int) -> bool:
    projects = get_projects(data)
    if not projects:
        return False

    section_heading(doc, "Projects")
    for i, project in enumerate(projects):
        project_header(doc, project.get("name", ""), get_project_tech(project), project.get("github_url", ""))
        for b in project.get("bullets") or []:
            bullet(doc, b, bold_markers)
        if i < len(projects) - 1:
            gap(doc, sub_gap)

    return True


# ── Build DOCX ───────────────────────────────────────────────────────────────
def build_docx(json_path: str, company_override: str = "", section_gap: int = DEFAULT_SECTION_GAP, sub_gap: int = DEFAULT_SUB_GAP) -> None:
    data = load_json(json_path)
    cfg = config(data)

    rtype = normalize_type(cfg.get("type", "fullstack"))
    level = normalize_level(cfg.get("level", 3))
    layout_profile = normalize_layout_profile(cfg.get("layout_profile", ""), rtype, level)

    if (rtype, level) not in CONFIGS:
        fail(f"Unsupported type/level: {rtype}, {level}")

    conf = CONFIGS[(rtype, level)]
    section_order = ["summary", "education", "experience", "projects", "skills"]
    company = company_override or clean(cfg.get("company", ""))

    if company_override:
        output_name = f"{RESUME_STEM}_{safe_name(company_override)}_Resume.docx"
    else:
        output_name = clean(cfg.get("output", "")) or f"{RESUME_STEM}_{safe_name(company or rtype)}_Resume.docx"

    if not output_name.lower().endswith(".docx"):
        output_name += ".docx"

    os.makedirs(WORD_DIR, exist_ok=True)
    output = os.path.join(WORD_DIR, os.path.basename(output_name))

    bold_markers = bool_value(cfg.get("bold_markers"), default=True)

    if os.path.exists(output):
        os.remove(output)

    doc = document_from_template()

    props = doc.core_properties
    props.author = clean(data.get("name", ""))
    props.title = f"{clean(data.get('name', ''))} {conf['label']}"
    props.subject = conf["label"]
    props.keywords = ATS_KEYWORDS.get(rtype, "")

    normal = doc.styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal.font.size = BODY_PT
    normal.paragraph_format.space_before = ZERO
    normal.paragraph_format.space_after = ZERO
    normal.paragraph_format.line_spacing = 1.0

    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p_name)
    r = p_name.add_run(clean(data.get("name", "")))
    rf(r, sz=NAME_PT, bold=True)

    contact = clean(data.get("contact", ""))
    raw_contact_lines = [line.strip() for line in contact.splitlines() if line.strip()] or [contact]
    contact_parts = []
    location_line = clean(data.get("location", "") or cfg.get("location", ""))

    for line in raw_contact_lines:
        parts = [p.strip() for p in line.split(" | ")] if " | " in line else [line]
        if parts and not location_line:
            first = parts[0]
            looks_like_location = "," in first and "@" not in first and "linkedin" not in first.lower() and "github" not in first.lower()
            if looks_like_location:
                location_line = first
                parts = parts[1:]
        contact_parts.extend(p for p in parts if p)

    if contact_parts:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp(p_contact)

        for i, part in enumerate(contact_parts):
            low = part.lower()
            if "linkedin" in low:
                hyperlink(p_contact, part, data.get("linkedin_url", ""), size=CONTACT_PT)
            elif "github" in low:
                hyperlink(p_contact, part, data.get("github_url", ""), size=CONTACT_PT)
            else:
                rr = p_contact.add_run(part)
                rf(rr, sz=CONTACT_PT)

            if i < len(contact_parts) - 1:
                sep = p_contact.add_run(" | ")
                rf(sep, sz=CONTACT_PT)

    citizenship = clean(data.get("citizenship", "") or cfg.get("citizenship", ""))
    if location_line or citizenship:
        p_location = doc.add_paragraph()
        p_location.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp(p_location)
        location_text = f"{citizenship} in {location_line}" if citizenship and location_line else citizenship or location_line
        rr = p_location.add_run(location_text)
        rf(rr, sz=CONTACT_PT)

    gap(doc, 6)

    renderers = {
        "summary":    lambda: render_summary(doc, data, bold_markers, sub_gap),
        "education":  lambda: render_education(doc, data, conf["grad"], level, bold_markers, sub_gap),
        "skills":     lambda: render_skills(doc, data, bold_markers, sub_gap),
        "experience": lambda: render_experience(doc, data, level, layout_profile, bold_markers, sub_gap),
        "projects":   lambda: render_projects(doc, data, bold_markers, sub_gap),
    }

    for section in section_order:
        rendered = renderers[section]()
        if rendered:
            gap(doc, section_gap)

    remove_final_empty_gap(doc)

    doc.save(output)

    print(f"DOCX saved   : {output}")
    print(f"Type         : {rtype}")
    print(f"Level        : {level} ({'Entry/SWE I' if level == 2 else 'Mid' if level == 3 else 'Intern'})")
    print(f"Layout profile: {layout_profile}")
    print(f"Section gap  : {section_gap}pt")
    print(f"Sub-section gap: {sub_gap}pt")
    print(f"Experience order: {experience_order_label(data, level, layout_profile)}")


# ── Main ─────────────────────────────────────────────────────────────────────
def main() -> None:
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(0)

    target = sys.argv[1]

    if target.lower() == "clear":
        clear_folder()
        return

    if target.lower().endswith(".json"):
        company, section_gap, sub_gap = parse_extra_args(sys.argv[2:])
        build_docx(target, company, section_gap, sub_gap)
        return

    if target.lower().endswith(".docx"):
        # For DOCX → PDF, only the company name matters. Gap args are ignored.
        company, _, _ = parse_extra_args(sys.argv[2:])
        docx_to_pdf(target, company)
        return

    fail("Use a .json file, a .docx file, or clear")


if __name__ == "__main__":
    main()
